import os
import tempfile
from typing import Any

import requests

from crewai_tools.rag.base_loader import BaseLoader, LoaderResult
from crewai_tools.rag.source_content import SourceContent


class DOCXLoader(BaseLoader):
    def load(self, source_content: SourceContent, **kwargs) -> LoaderResult:  # type: ignore[override]
        try:
            from docx import Document as DocxDocument
        except ImportError as e:
            raise ImportError(
                "python-docx is required for DOCX loading. Install with: 'uv pip install python-docx' or pip install crewai-tools[rag]"
            ) from e

        source_ref = source_content.source_ref

        if source_content.is_url():
            temp_file = self._download_from_url(source_ref, kwargs)
            try:
                return self._load_from_file(temp_file, source_ref, DocxDocument)
            finally:
                os.unlink(temp_file)
        elif source_content.path_exists():
            return self._load_from_file(source_ref, source_ref, DocxDocument)
        else:
            raise ValueError(
                f"Source must be a valid file path or URL, got: {source_content.source}"
            )

    @staticmethod
    def _download_from_url(url: str, kwargs: dict) -> str:
        headers = kwargs.get(
            "headers",
            {
                "Accept": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "User-Agent": "Mozilla/5.0 (compatible; crewai-tools DOCXLoader)",
            },
        )

        try:
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()

            # Create temporary file to save the DOCX content
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_file.write(response.content)
                return temp_file.name
        except Exception as e:
            raise ValueError(f"Error fetching content from URL {url}: {e!s}") from e

    def _load_from_file(
        self,
        file_path: str,
        source_ref: str,
        DocxDocument: Any,  # noqa: N803
    ) -> LoaderResult:
        try:
            doc = DocxDocument(file_path)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)  # noqa: PERF401

            content = "\n".join(text_parts)

            metadata = {
                "format": "docx",
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }

            return LoaderResult(
                content=content,
                source=source_ref,
                metadata=metadata,
                doc_id=self.generate_doc_id(source_ref=source_ref, content=content),
            )

        except Exception as e:
            raise ValueError(f"Error loading DOCX file: {e!s}") from e
